#!/usr/bin/env python3

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os
import shutil
from datetime import datetime
from pathlib import Path

class MSADocumentBuilder:
    def __init__(self):
        self.output_dir = 'output'
        # Use multiple fallback approaches to find the template file in deployment environments
        self.template_path = self._find_template_file()
        os.makedirs(self.output_dir, exist_ok=True)
    
    def _find_template_file(self):
        """Find the template file using multiple fallback approaches for deployment environments"""
        template_filename = 'original_template.docx'
        
        # Method 1: Same directory as this script (normal case)
        script_dir_path = Path(__file__).parent / template_filename
        if script_dir_path.exists():
            print(f"Template found via script directory: {script_dir_path}")
            return script_dir_path
        
        # Method 2: Current working directory (deployment case)
        cwd_path = Path.cwd() / template_filename
        if cwd_path.exists():
            print(f"Template found via current working directory: {cwd_path}")
            return cwd_path
        
        # Method 3: Root directory (some deployment environments)
        root_path = Path('/') / template_filename
        if root_path.exists():
            print(f"Template found via root directory: {root_path}")
            return root_path
        
        # Method 4: App directory (common in cloud deployments)
        app_path = Path('/app') / template_filename
        if app_path.exists():
            print(f"Template found via /app directory: {app_path}")
            return app_path
        
        # Method 5: Relative to current working directory with different depths
        for depth in ['', './', '../', '../../']:
            candidate_path = Path(depth) / template_filename
            if candidate_path.exists():
                print(f"Template found via relative path {depth}: {candidate_path.resolve()}")
                return candidate_path.resolve()
        
        # Method 6: Search in common deployment paths
        common_paths = [
            '/var/app/current',
            '/app',
            '/home/app',
            '/usr/src/app',
            os.path.expanduser('~')
        ]
        
        for base_path in common_paths:
            candidate_path = Path(base_path) / template_filename
            if candidate_path.exists():
                print(f"Template found via common path {base_path}: {candidate_path}")
                return candidate_path
        
        # If all else fails, use the original approach (script directory)
        fallback_path = Path(__file__).parent / template_filename
        print(f"Using fallback path (may not exist): {fallback_path}")
        return fallback_path
    
    def generate_msa(self, client_data, preparer_data, include_compliance, include_security_plus, pricing_model, pricing_data):
        """Generate a complete MSA document using the original template as base"""
        
        print("Generating MSA using original template...")
        print(f"Client: {client_data['name']}")
        print(f"Preparer: {preparer_data['name']}")
        print(f"Pricing model: {pricing_model}")
        print(f"Include compliance: {include_compliance}")
        print(f"Include security plus: {include_security_plus}")
        print(f"Template path being used: {self.template_path}")
        
        try:
            # Load the original template document
            if not self.template_path.exists():
                # Additional debugging information
                print(f"Template file not found at: {self.template_path}")
                print(f"Current working directory: {Path.cwd()}")
                print(f"Script file location: {Path(__file__).parent}")
                print("Files in current directory:", list(Path.cwd().iterdir()))
                print("Files in script directory:", list(Path(__file__).parent.iterdir()))
                raise Exception(f"Template file not found: {self.template_path}")
            
            doc = Document(str(self.template_path))
            
            # 1. Add client information sections in the lower portion of page 1
            self._add_client_sections(doc, client_data, preparer_data)
            
            # 2. Replace client-specific information in document body
            self._replace_client_information(doc, client_data)
            
            # 3. Replace dates
            self._replace_dates(doc)
            
            # 4. Handle optional services (add/remove sections based on selections)
            self._handle_optional_services(doc, include_compliance, include_security_plus)
            
            # 5. Update pricing table with dynamic pricing
            self._update_pricing_table(doc, pricing_model, pricing_data, include_compliance, include_security_plus)
            
            # 6. Save the generated document
            filename = f"MSA_{client_data['name'].replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            output_path = os.path.join(self.output_dir, filename)
            doc.save(output_path)
            
            print(f"MSA document generated: {output_path}")
            return output_path
            
        except Exception as e:
            print(f"Error generating MSA document: {e}")
            raise e
    
    def _add_client_sections(self, doc, client_data, preparer_data):
        """Add client information sections on page 1 after the introduction"""
        
        print("Adding client sections on page 1...")
        
        # Find the insertion point early in the document (page 1)
        insertion_point = None
        
        # Look for "Our Core Values" section to insert client info before it
        for i, paragraph in enumerate(doc.paragraphs[:50]):  # Only search first 50 paragraphs
            text = paragraph.text.strip()
            if text == "Our Core Values":
                insertion_point = i
                break
        
        # Fallback: if we can't find "Our Core Values", insert after the main introduction
        # (after the Journey to IT Maturity paragraph)
        if insertion_point is None:
            for i, paragraph in enumerate(doc.paragraphs[:20]):  # Search first 20 paragraphs
                text = paragraph.text.strip().lower()
                if "journey to it maturity" in text or "it maturity" in text:
                    insertion_point = i + 1  # Insert right after this paragraph
                    break
        
        # Final fallback: insert after paragraph 6 (after main intro content)
        if insertion_point is None:
            insertion_point = 7
        
        print(f"Inserting client sections at paragraph index {insertion_point}")
        
        # Get the paragraph element where we'll insert
        space_para = doc.paragraphs[insertion_point]._element
        
        # Create new paragraphs for client sections
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls, qn
        
        # Add minimal spacing before client sections (page 1 insertion)
        for _ in range(2):  # Reduced to 2 line breaks for page 1
            p_xml = f'<w:p {nsdecls("w")}><w:pPr/></w:p>'
            new_p = parse_xml(p_xml)
            space_para.getparent().insert(space_para.getparent().index(space_para) + 1, new_p)
        
        # Get the last spacer paragraph for further insertions
        current_para = space_para
        for _ in range(2):  # Reduced to 2 
            current_para = current_para.getnext()
        
        print("Client sections added successfully on page 1")
        
        # Create and insert paragraphs one by one
        paragraphs_to_add = []
        
        # "Prepared For:" section
        paragraphs_to_add.extend([
            f'<w:p {nsdecls("w")}><w:pPr><w:spacing w:after="60"/></w:pPr><w:r><w:rPr><w:b/></w:rPr><w:t>Prepared For:</w:t></w:r></w:p>',
            f'<w:p {nsdecls("w")}><w:pPr><w:spacing w:after="0"/></w:pPr><w:r><w:t>{client_data.get("name", "")}</w:t></w:r></w:p>',
            f'<w:p {nsdecls("w")}><w:pPr><w:spacing w:after="0"/></w:pPr><w:r><w:t>{client_data.get("email", "")}</w:t></w:r></w:p>'
        ])
        
        # Add address lines
        address = client_data.get('address', '')
        if address:
            address_lines = []
            if '\\n' in address:
                address_lines = address.split('\\n')
            elif ',' in address and len(address) > 50:
                parts = address.split(',')
                address_lines = [parts[0].strip()]
                if len(parts) > 1:
                    address_lines.append(','.join(parts[1:]).strip())
            else:
                address_lines = [address]
            
            for addr_line in address_lines:
                if addr_line.strip():
                    paragraphs_to_add.append(f'<w:p {nsdecls("w")}><w:pPr><w:spacing w:after="0"/></w:pPr><w:r><w:t>{addr_line.strip()}</w:t></w:r></w:p>')
        
        # Add phone if available
        phone = client_data.get('phone', '')
        if phone:
            paragraphs_to_add.append(f'<w:p {nsdecls("w")}><w:pPr><w:spacing w:after="120"/></w:pPr><w:r><w:t>{phone}</w:t></w:r></w:p>')
        else:
            paragraphs_to_add.append(f'<w:p {nsdecls("w")}><w:pPr><w:spacing w:after="120"/></w:pPr><w:r><w:t></w:t></w:r></w:p>')
        
        # "Prepared By:" section
        paragraphs_to_add.extend([
            f'<w:p {nsdecls("w")}><w:pPr><w:spacing w:after="60"/></w:pPr><w:r><w:rPr><w:b/></w:rPr><w:t>Prepared By:</w:t></w:r></w:p>',
            f'<w:p {nsdecls("w")}><w:pPr><w:spacing w:after="0"/></w:pPr><w:r><w:t>{preparer_data.get("name", "Kevin Fuller")}</w:t></w:r></w:p>',
            f'<w:p {nsdecls("w")}><w:pPr><w:spacing w:after="0"/></w:pPr><w:r><w:t>{preparer_data.get("email", "k.fuller@avatarmsp.com")}</w:t></w:r></w:p>',
            f'<w:p {nsdecls("w")}><w:pPr><w:spacing w:after="120"/></w:pPr><w:r><w:t></w:t></w:r></w:p>'
        ])
        
        # Insert paragraphs
        for para_xml in paragraphs_to_add:
            new_para = parse_xml(para_xml)
            current_para.getparent().insert(current_para.getparent().index(current_para) + 1, new_para)
            current_para = new_para
        
        print(f"Added {len(paragraphs_to_add)} client information paragraphs")
    
    def _replace_client_information(self, doc, client_data):
        """Replace placeholder client information with actual client data"""
        
        print("Replacing client information...")
        
        # Common replacements that might exist in the template
        replacements = {
            'Katy Spring Solutions': client_data['name'],
            'Katy Spring': client_data['name'],
            'CLIENT_NAME': client_data['name'],
            '{{client_name}}': client_data['name'],
        }
        
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            for old_text, new_text in replacements.items():
                if old_text in paragraph.text:
                    for run in paragraph.runs:
                        if old_text in run.text:
                            run.text = run.text.replace(old_text, new_text)
                            print(f"Replaced '{old_text}' with '{new_text}' in paragraph")
        
        # Replace in tables
        for table_idx, table in enumerate(doc.tables):
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for old_text, new_text in replacements.items():
                            if old_text in paragraph.text:
                                for run in paragraph.runs:
                                    if old_text in run.text:
                                        run.text = run.text.replace(old_text, new_text)
                                        print(f"Replaced '{old_text}' with '{new_text}' in table {table_idx}")
    
    def _replace_dates(self, doc):
        """Replace date placeholders with current dates"""
        
        print("Replacing dates...")
        
        current_date = datetime.now()
        current_month = current_date.strftime('%B')
        current_year = str(current_date.year)
        
        date_replacements = {
            'July 2025': f'{current_month} {current_year}',
            'DATE_PLACEHOLDER': current_date.strftime('%B %d, %Y'),
            '{{current_date}}': current_date.strftime('%B %d, %Y'),
            '{{current_month_year}}': f'{current_month} {current_year}'
        }
        
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            for old_text, new_text in date_replacements.items():
                if old_text in paragraph.text:
                    for run in paragraph.runs:
                        if old_text in run.text:
                            run.text = run.text.replace(old_text, new_text)
                            print(f"Replaced date '{old_text}' with '{new_text}'")
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for old_text, new_text in date_replacements.items():
                            if old_text in paragraph.text:
                                for run in paragraph.runs:
                                    if old_text in run.text:
                                        run.text = run.text.replace(old_text, new_text)
                                        print(f"Replaced date '{old_text}' with '{new_text}' in table")
    
    def _handle_optional_services(self, doc, include_compliance, include_security_plus):
        """Handle optional service sections based on user selections"""
        
        print(f"Handling optional services: compliance={include_compliance}, security_plus={include_security_plus}")
        
        # This would need specific implementation based on your template structure
        # For now, we'll just log what we would do
        
        if include_compliance:
            print("Including compliance services section")
        else:
            print("Excluding compliance services section")
            
        if include_security_plus:
            print("Including security plus services section")
        else:
            print("Excluding security plus services section")
        
        # TODO: Implement actual section addition/removal logic based on template structure
    
    def _update_pricing_table(self, doc, pricing_model, pricing_data, include_compliance, include_security_plus):
        """Update the pricing table with dynamic pricing based on selected model"""
        
        print(f"Updating pricing table for model: {pricing_model}")
        print(f"Pricing data: {pricing_data}")
        
        # Find and update the pricing table (usually the first table)
        if not doc.tables:
            print("No tables found in document")
            return
        
        pricing_table = doc.tables[0]  # Assuming the first table is the pricing table
        
        try:
            if pricing_model == 'workstation':
                self._update_workstation_pricing(pricing_table, pricing_data, include_compliance, include_security_plus)
            else:  # user model
                self._update_user_pricing(pricing_table, pricing_data, include_compliance, include_security_plus)
        except Exception as e:
            print(f"Error updating pricing table: {e}")
            # Continue without failing the entire document generation
    
    def _update_workstation_pricing(self, table, pricing_data, include_compliance, include_security_plus):
        """Update pricing table for workstation-based model"""
        
        print("Updating workstation pricing...")
        
        workstation_count = pricing_data.get('workstation_count', 0)
        workstation_price = pricing_data.get('workstation_price', 110.00)
        
        # Calculate base cost
        base_monthly_cost = workstation_count * workstation_price
        
        # Add optional services
        compliance_cost = 0
        security_plus_cost = 0
        
        if include_compliance:
            compliance_cost = workstation_count * 10.00  # $10 per workstation
        
        if include_security_plus:
            security_plus_cost = workstation_count * 15.00  # $15 per workstation
        
        total_monthly_cost = base_monthly_cost + compliance_cost + security_plus_cost
        
        # Update table rows
        # This assumes a specific table structure - adjust based on your actual template
        try:
            # Example table structure update
            if len(table.rows) >= 2:
                # Update quantity
                if len(table.rows[1].cells) >= 2:
                    table.rows[1].cells[1].text = str(workstation_count)
                
                # Update unit price
                if len(table.rows[1].cells) >= 3:
                    table.rows[1].cells[2].text = f"${workstation_price:.2f}"
                
                # Update total
                if len(table.rows[1].cells) >= 4:
                    table.rows[1].cells[3].text = f"${base_monthly_cost:.2f}"
                    
            print(f"Updated workstation pricing: {workstation_count} workstations x ${workstation_price:.2f} = ${total_monthly_cost:.2f}/month")
            
        except Exception as e:
            print(f"Error updating workstation pricing table: {e}")
    
    def _update_user_pricing(self, table, pricing_data, include_compliance, include_security_plus):
        """Update pricing table for user-based model"""
        
        print("Updating user pricing...")
        
        user_count = pricing_data.get('user_count', 0)
        user_price = pricing_data.get('user_price', 15.00)
        
        # Calculate base cost
        base_monthly_cost = user_count * user_price
        
        # Add optional services
        compliance_cost = 0
        security_plus_cost = 0
        
        if include_compliance:
            compliance_cost = user_count * 5.00  # $5 per user
        
        if include_security_plus:
            security_plus_cost = user_count * 8.00  # $8 per user
        
        total_monthly_cost = base_monthly_cost + compliance_cost + security_plus_cost
        
        # Update table rows
        try:
            # Example table structure update
            if len(table.rows) >= 2:
                # Update quantity
                if len(table.rows[1].cells) >= 2:
                    table.rows[1].cells[1].text = str(user_count)
                
                # Update unit price
                if len(table.rows[1].cells) >= 3:
                    table.rows[1].cells[2].text = f"${user_price:.2f}"
                
                # Update total
                if len(table.rows[1].cells) >= 4:
                    table.rows[1].cells[3].text = f"${base_monthly_cost:.2f}"
                    
            print(f"Updated user pricing: {user_count} users x ${user_price:.2f} = ${total_monthly_cost:.2f}/month")
            
        except Exception as e:
            print(f"Error updating user pricing table: {e}")

# For testing
if __name__ == '__main__':
    # Test data
    test_client_data = {
        'name': 'Test Company Inc.',
        'email': 'contact@testcompany.com',
        'address': '123 Test Street\\nSuite 456\\nTest City, TX 12345',
        'phone': '(555) 123-4567'
    }
    
    test_preparer_data = {
        'name': 'Kevin Fuller',
        'email': 'k.fuller@avatarmsp.com'
    }
    
    test_pricing_data = {
        'workstation_count': 25,
        'workstation_price': 110.00
    }
    
    builder = MSADocumentBuilder()
    
    try:
        output_path = builder.generate_msa(
            client_data=test_client_data,
            preparer_data=test_preparer_data,
            include_compliance=True,
            include_security_plus=False,
            pricing_model='workstation',
            pricing_data=test_pricing_data
        )
        print(f"Test MSA generated successfully: {output_path}")
    except Exception as e:
        print(f"Error generating test MSA: {e}")
